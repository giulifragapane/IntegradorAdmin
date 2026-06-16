"""Convert GUIA_VIDEO_ADMIN.md to Word (.docx)."""
from pathlib import Path
import re

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_shading(cell, fill: str):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shading)


def add_formatted_paragraph(doc, text: str, style=None, quote=False):
    """Parse **bold**, *italic*, `code` into runs."""
    p = doc.add_paragraph(style=style)
    if quote:
        p.paragraph_format.left_indent = Inches(0.4)
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)

    pattern = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`|\[[^\]]+\])")
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            p.add_run(text[pos : m.start()])
        chunk = m.group(0)
        if chunk.startswith("**") and chunk.endswith("**"):
            run = p.add_run(chunk[2:-2])
            run.bold = True
        elif chunk.startswith("*") and chunk.endswith("*"):
            run = p.add_run(chunk[1:-1])
            run.italic = True
        elif chunk.startswith("`") and chunk.endswith("`"):
            run = p.add_run(chunk[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x80, 0x00, 0x00)
        else:
            p.add_run(chunk)
        pos = m.end()
    if pos < len(text):
        p.add_run(text[pos:])
    return p


def parse_table_row(line: str) -> list[str]:
    line = line.strip()
    if line.startswith("|"):
        line = line[1:]
    if line.endswith("|"):
        line = line[:-1]
    return [c.strip() for c in line.split("|")]


def is_separator_row(cells: list[str]) -> bool:
    return all(re.match(r"^:?-+:?$", c.replace(" ", "")) or c == "" for c in cells)


def convert_md_to_docx(md_path: Path, docx_path: Path):
    lines = md_path.read_text(encoding="utf-8").splitlines()
    doc = Document()

    # Title
    title = doc.add_heading("Guion — Frontend de Administración (admin-app)", 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    i = 0
    while i < len(lines):
        line = lines[i]

        # Skip duplicate top title
        if line.startswith("# Guion"):
            i += 1
            continue

        if line.strip() == "---":
            i += 1
            continue

        # Headings
        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=1)
            i += 1
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=2)
            i += 1
            continue

        # Blockquote
        if line.startswith("> "):
            text = line[2:].strip()
            while i + 1 < len(lines) and lines[i + 1].startswith("> "):
                i += 1
                text += " " + lines[i][2:].strip()
            add_formatted_paragraph(doc, text, quote=True)
            i += 1
            continue

        # Table
        if line.strip().startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            rows = [parse_table_row(l) for l in table_lines]
            rows = [r for r in rows if not is_separator_row(r)]
            if rows:
                tbl = doc.add_table(rows=len(rows), cols=len(rows[0]))
                tbl.style = "Table Grid"
                for ri, row in enumerate(rows):
                    for ci, cell_text in enumerate(row):
                        cell = tbl.rows[ri].cells[ci]
                        cell.text = ""
                        add_formatted_paragraph(
                            Document(cell._tc), cell_text
                        )  # wrong approach
                # Fix: set cell text simply
                for ri, row in enumerate(rows):
                    for ci, cell_text in enumerate(row):
                        cell = tbl.rows[ri].cells[ci]
                        cell.text = re.sub(r"\*\*([^*]+)\*\*", r"\1", cell_text)
                        cell.text = re.sub(r"`([^`]+)`", r"\1", cell.text)
                        if ri == 0:
                            set_cell_shading(cell, "D9E2F3")
                            for p in cell.paragraphs:
                                for run in p.runs:
                                    run.bold = True
                doc.add_paragraph()
            continue

        # Checklist
        if line.strip().startswith("- [ ]") or line.strip().startswith("- [x]"):
            checked = "[x]" in line[:6].lower() if len(line) > 5 else False
            text = re.sub(r"^- \[[ xX]\]\s*", "", line.strip())
            p = doc.add_paragraph(style="List Bullet")
            sym = "☑ " if checked else "☐ "
            add_formatted_paragraph(doc, sym + text)
            # remove duplicate empty - use single paragraph
            doc.paragraphs[-1]._element.getparent().remove(doc.paragraphs[-1]._element)
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(sym)
            pos = 0
            for m in re.finditer(r"(\*\*[^*]+\*\*|`[^`]+`)", text):
                if m.start() > pos:
                    p.add_run(text[pos : m.start()])
                chunk = m.group(0)
                if chunk.startswith("**"):
                    r = p.add_run(chunk[2:-2])
                    r.bold = True
                else:
                    r = p.add_run(chunk[1:-1])
                    r.font.name = "Consolas"
                pos = m.end()
            if pos < len(text):
                p.add_run(text[pos:])
            i += 1
            continue

        # Bullet list
        if line.strip().startswith("- ") and not line.strip().startswith("- ["):
            text = line.strip()[2:]
            p = doc.add_paragraph(style="List Bullet")
            add_formatted_to_paragraph(p, text)
            i += 1
            continue

        # Empty
        if not line.strip():
            i += 1
            continue

        # Label lines like **Abrir:**
        if line.strip().startswith("**") and "**:" in line:
            add_formatted_paragraph(doc, line.strip())
            i += 1
            continue

        # Regular paragraph
        add_formatted_paragraph(doc, line.strip())
        i += 1

    doc.save(docx_path)
    print(f"Created: {docx_path}")


def add_formatted_to_paragraph(p, text: str):
    pattern = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`)")
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            p.add_run(text[pos : m.start()])
        chunk = m.group(0)
        if chunk.startswith("**"):
            r = p.add_run(chunk[2:-2])
            r.bold = True
        else:
            r = p.add_run(chunk[1:-1])
            r.font.name = "Consolas"
            r.font.size = Pt(10)
        pos = m.end()
    if pos < len(text):
        p.add_run(text[pos:])


# Fix table cell formatting - rewrite convert with cleaner table handling
def convert_md_to_docx_v2(md_path: Path, docx_path: Path):
    lines = md_path.read_text(encoding="utf-8").splitlines()
    doc = Document()

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    h = doc.add_heading("Guion — Frontend de Administración (admin-app)", 0)
    h.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    i = 0
    while i < len(lines):
        line = lines[i]

        if line.startswith("# Guion") or line.strip() == "---":
            i += 1
            continue

        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=1)
            i += 1
            continue

        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=2)
            i += 1
            continue

        if line.startswith("> "):
            parts = []
            while i < len(lines) and lines[i].startswith("> "):
                parts.append(lines[i][2:].strip())
                i += 1
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.space_after = Pt(8)
            add_formatted_to_paragraph(p, " ".join(parts))
            continue

        if line.strip().startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            rows = [parse_table_row(l) for l in table_lines]
            rows = [r for r in rows if not is_separator_row(r)]
            if rows:
                ncols = max(len(r) for r in rows)
                tbl = doc.add_table(rows=len(rows), cols=ncols)
                tbl.style = "Table Grid"
                for ri, row in enumerate(rows):
                    for ci in range(ncols):
                        txt = row[ci] if ci < len(row) else ""
                        txt = re.sub(r"\*\*([^*]+)\*\*", r"\1", txt)
                        txt = re.sub(r"`([^`]+)`", r"\1", txt)
                        cell = tbl.rows[ri].cells[ci]
                        cell.text = txt
                        if ri == 0:
                            set_cell_shading(cell, "D9E2F3")
                            for p in cell.paragraphs:
                                for run in p.runs:
                                    run.bold = True
                doc.add_paragraph()
            continue

        if re.match(r"^- \[[ xX]\]", line.strip()):
            checked = bool(re.match(r"^- \[[xX]\]", line.strip()))
            text = re.sub(r"^- \[[ xX]\]\s*", "", line.strip())
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(("☑ " if checked else "☐ "))
            add_formatted_to_paragraph(p, text)
            i += 1
            continue

        if line.strip().startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_formatted_to_paragraph(p, line.strip()[2:])
            i += 1
            continue

        if not line.strip():
            i += 1
            continue

        p = doc.add_paragraph()
        add_formatted_to_paragraph(p, line.strip())
        i += 1

    doc.save(docx_path)
    print(f"Created: {docx_path}")


if __name__ == "__main__":
    root = Path(__file__).resolve().parent.parent
    md = root / "GUIA_VIDEO_ADMIN.md"
    out = root / "GUIA_VIDEO_ADMIN.docx"
    convert_md_to_docx_v2(md, out)
